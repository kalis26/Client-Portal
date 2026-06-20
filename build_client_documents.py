from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path('client-packet-overhaul')
OUT.mkdir(exist_ok=True)

NAVY='0B2545'; BLUE='2E74B5'; DARK='1F4D78'; MUTED='5B6573'; LIGHT='F2F4F7'; PALE='E8EEF5'; GOLD='7A5A00'
CONTACT='Rachid Mustapha Amine | Full-Stack Web Developer\naminera2006@gmail.com | +213 798 19 92 08\nGitHub: github.com/kalis26 | LinkedIn: linkedin.com/in/amine-mustapha-rachid'

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); mar = tcPr.first_child_found_in('w:tcMar')
    if mar is None: mar = OxmlElement('w:tcMar'); tcPr.append(mar)
    for side, value in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = mar.find(qn(f'w:{side}'))
        if node is None: node=OxmlElement(f'w:{side}'); mar.append(node)
        node.set(qn('w:w'), str(value)); node.set(qn('w:type'),'dxa')

def borders(table, color='D9E0E8'):
    tblPr=table._tbl.tblPr; el=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        x=OxmlElement(f'w:{edge}'); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),'6'); x.set(qn('w:color'),color); el.append(x)
    tblPr.append(el)

def set_widths(table, widths):
    table.autofit=False
    for row in table.rows:
        for cell, width in zip(row.cells,widths): cell.width=Inches(width)

def font(run, size=11, color=NAVY, bold=False, italic=False):
    run.font.name='Aptos'; run._element.rPr.rFonts.set(qn('w:ascii'),'Aptos'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Aptos')
    run.font.size=Pt(size); run.font.color.rgb=RGBColor.from_string(color); run.bold=bold; run.italic=italic

def base_doc(label):
    d=Document(); sec=d.sections[0]
    sec.top_margin=Inches(.72); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.78); sec.right_margin=Inches(.78)
    sec.header_distance=Inches(.32); sec.footer_distance=Inches(.32)
    normal=d.styles['Normal']; normal.font.name='Aptos'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Aptos'); normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
    for name,size,color,before,after in [('Title',28,NAVY,0,6),('Heading 1',16,BLUE,16,7),('Heading 2',12.5,DARK,10,4)]:
        s=d.styles[name]; s.font.name='Aptos Display' if name=='Title' else 'Aptos'; s._element.rPr.rFonts.set(qn('w:ascii'),s.font.name); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=header.add_run(label.upper()); font(r,8.5,MUTED,True)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=footer.add_run('Rachid Mustapha Amine  |  Full-Stack Web Developer'); font(r,8,MUTED)
    return d

def p(d, text='', style=None, bold_prefix=None, color=NAVY, size=10.5):
    para=d.add_paragraph(style=style) if style else d.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r=para.add_run(bold_prefix); font(r,size,color,True); r=para.add_run(text[len(bold_prefix):]); font(r,size,color)
    else:
        r=para.add_run(text); font(r,size,color)
    return para

def bullets(d, items):
    for item in items:
        para=d.add_paragraph(style='List Bullet'); para.paragraph_format.space_after=Pt(3); r=para.add_run(item); font(r,10.3,NAVY)

def title(d, kicker, heading, sub=''):
    p1=d.add_paragraph(); p1.paragraph_format.space_after=Pt(2); r=p1.add_run(kicker.upper()); font(r,9,GOLD,True)
    h=d.add_paragraph(style='Title'); r=h.add_run(heading); font(r,28,NAVY,True)
    if sub: p(d,sub,color=MUTED,size=12)

def meta(d, rows):
    t=d.add_table(rows=0, cols=2); t.alignment=WD_TABLE_ALIGNMENT.LEFT; set_widths(t,[1.35,5.1]); borders(t)
    for label,value in rows:
        cells=t.add_row().cells
        for c in cells: cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cells[0],LIGHT); rr=cells[0].paragraphs[0].add_run(label); font(rr,9.5,DARK,True)
        rr=cells[1].paragraphs[0].add_run(value); font(rr,9.5,NAVY)
    return t

def note(d, heading, text):
    t=d.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.LEFT; set_widths(t,[6.45]); borders(t,'C8D5E3'); cell=t.cell(0,0); shade(cell,PALE); cell_margins(cell,140,180,140,180)
    para=cell.paragraphs[0]; r=para.add_run(heading+' '); font(r,10,DARK,True); r=para.add_run(text); font(r,10,NAVY)

def onboarding():
    d=base_doc('Client onboarding guide')
    title(d,'Welcome','Your website or web app project starts here','A practical guide to kickoff, collaboration, launch, and handover.')
    meta(d,[('Prepared for:','[Client / Company Name]'),('Project:','[Project name]'),('Project workspace:','[Client portal / shared workspace link]'),('Primary contact:','Rachid Mustapha Amine')])
    p(d,'Welcome',style='Heading 1')
    p(d,'Thank you for choosing to work with me. I build responsive business websites and custom web applications that are clear to use, reliable to operate, and ready to grow with your business.')
    p(d,'What I build',style='Heading 1')
    bullets(d,['Business websites and landing pages that communicate your offer clearly.','E-commerce experiences with product, cart, account, and order flows.','Dashboards and internal tools that make day-to-day operations easier.','Custom web apps with secure authentication, APIs, databases, and deployment.'])
    note(d,'How this is delivered:','Each project is scoped individually. Design, frontend, backend/API work, authentication, PostgreSQL data storage, integrations, and launch support are included only where confirmed in your proposal.')
    p(d,'How we will work',style='Heading 1')
    steps=[('1. Kickoff','We confirm your goals, audience, priorities, scope, and access requirements.'),('2. Structure & design','We agree on page structure, user flows, content priorities, and visual direction.'),('3. Build','I implement the approved scope and share progress through the agreed workspace.'),('4. Review','You provide one consolidated set of feedback; two revision rounds are included unless the proposal says otherwise.'),('5. Launch & handover','After final payment, I deliver the agreed assets and support the agreed launch process.')]
    for h,txt in steps:
        pp=d.add_paragraph(); pp.paragraph_format.space_after=Pt(5); r=pp.add_run(h+'  '); font(r,10.8,BLUE,True); r=pp.add_run(txt); font(r,10.5,NAVY)
    p(d,'What I need from you',style='Heading 1')
    bullets(d,['A clear business goal and the priority pages, features, or customer actions.','Brand assets, approved copy, images, product/service details, and legal text where applicable.','Access to your domain, hosting, analytics, payment provider, or third-party tools when required.','One decision-maker or a single consolidated feedback response for each review cycle.'])
    p(d,'Communication and next actions',style='Heading 1')
    bullets(d,['Primary channels: email and WhatsApp. Typical response target: within one business day.','Use the project workspace for files, progress updates, approvals, and final links.','To begin: sign the project agreement, pay the 20% deposit, and complete the kickoff information above.'])
    note(d,'After launch:','The agreed handover may include source code, design files, credentials transfer, deployment notes, or short training. Ongoing maintenance, new features, and support are separate services unless included in the proposal.')
    p(d,CONTACT,color=MUTED,size=9.4)
    d.save(OUT/'Client_Onboarding_Guide.docx')

def agreement():
    d=base_doc('Service agreement template')
    title(d,'Service agreement','Website & web application services','Operational template - obtain Algerian legal review before using this as a binding standard agreement.')
    meta(d,[('Client:','[Client legal name / company]'),('Project:','[Project name]'),('Proposal / Scope:','[Proposal reference and date]'),('Effective date:','[Date signed by both parties]')])
    note(d,'Purpose:','This agreement sets the general terms for the project. The signed proposal or statement of work controls the specific deliverables, exclusions, timeline, milestones, fees, and acceptance criteria.')
    sections=[
    ('1. Services and scope','The Developer will provide the services described in the attached or referenced proposal. Requests that are not included in that scope, including new pages, features, integrations, or changed requirements, require written approval and may require a revised fee and timeline.'),
    ('2. Start date and timeline','Work begins only after both parties have signed this agreement and the 20% upfront payment has cleared. Any delivery dates are estimates unless the proposal expressly states otherwise. Client delays, missing content, late approvals, or unavailable access extend the timeline by a reasonable corresponding period.'),
    ('3. Fees and payment','The Client will pay 20% of the agreed project fee before work begins and the remaining 80% before final handover, transfer of source files, or production launch, unless the proposal states another schedule. Invoices are payable in the single currency shown on that invoice. Payment methods and due dates are shown on the invoice.'),
    ('4. Revisions and approvals','The project includes two rounds of revisions per deliverable unless the proposal states otherwise. A revision means adjustments to an approved direction or agreed deliverable. A new concept, feature, page, workflow, or substantial change in direction is a change request. The Client must provide clear, consolidated feedback within the agreed review period.'),
    ('5. Client responsibilities','The Client will provide accurate content, approvals, decisions, brand assets, and timely access to relevant accounts or systems. The Client confirms it has the rights to use all material it supplies, including text, images, trademarks, data, and media.'),
    ('6. Change requests','Either party may identify work outside the agreed scope. The Developer will provide the expected impact on fees and timeline before starting it. The Developer is not required to begin additional work until the change is approved in writing.'),
    ('7. Third-party services and access','Domains, hosting, payment processors, email services, fonts, stock assets, APIs, plugins, and similar services may be subject to separate terms and fees. Unless the proposal states otherwise, the Client owns and pays for those accounts. The Client is responsible for maintaining necessary access and licenses.'),
    ('8. Delivery, launch, and handover','Final deliverables are those listed in the proposal. Launch depends on the Client providing required approvals, credentials, domain/hosting access, and legal or business content. Source files, code, credentials transfer, and final deployment access are provided only after full payment and only where included in the proposal.'),
    ('9. Ownership and portfolio use','Once all project fees are paid, the Client receives ownership or the license to use the final deliverables described in the proposal. The Developer retains ownership of pre-existing tools, reusable code, methods, and third-party materials, subject to their licenses. The Developer may display completed work in a portfolio unless the parties agree otherwise in writing.'),
    ('10. Confidentiality','Each party will use the other party\'s non-public business, technical, and customer information only for this project and will not disclose it except where needed to perform the services or where legally required.'),
    ('11. Cancellation and suspension','Either party may cancel the project with written notice. The Client must pay for work completed and approved expenses up to the cancellation date. The initial payment is non-refundable because it reserves project capacity and covers initial work. The Developer may pause or terminate work for non-payment, prolonged Client delay, or abusive conduct.'),
    ('12. Liability and warranty','The Developer will perform the services with reasonable professional care. To the maximum extent allowed by applicable law, the Developer is not liable for indirect or consequential losses, third-party service failures, or losses arising from Client-provided content or instructions. Any specific post-launch support or warranty period must be stated in the proposal.'),
    ('13. General terms','Written approval includes email or the agreed project workspace. This agreement and the proposal form the complete agreement for the project. Any amendment must be in writing and accepted by both parties. The parties should seek Algerian legal review before relying on this template as a binding standard agreement.')]
    for h,body in sections:
        p(d,h,style='Heading 2'); p(d,body)
    p(d,'Signatures',style='Heading 1')
    t=d.add_table(rows=4,cols=2); set_widths(t,[3.2,3.2]); borders(t)
    labels=[('Client name / title:','Developer: Rachid Mustapha Amine'),('Signature:','Signature:'),('Date:','Date:'),('Company / address:','Address: Ouled Chebel, Alger')]
    for row,vals in zip(t.rows,labels):
        for c,v in zip(row.cells,vals): cell_margins(c,160,140,300,140); r=c.paragraphs[0].add_run(v); font(r,10,NAVY)
    d.save(OUT/'Website_Web_App_Service_Agreement.docx')

def invoice():
    d=base_doc('Invoice template')
    title(d,'Invoice','[Invoice number]','Website & web application services')
    meta(d,[('Issue date:','[DD Month YYYY]'),('Due date:','[DD Month YYYY]'),('Currency:','[DZD / EUR / USD - choose one]'),('Payment status:','[Unpaid / Partially paid / Paid]')])
    p(d,'Bill to',style='Heading 1')
    meta(d,[('Client / company:','[Legal name]'),('Address:','[Address]'),('Email:','[Email]'),('Phone / tax ID:','[Phone / tax ID if applicable]')])
    p(d,'Project and services',style='Heading 1')
    t=d.add_table(rows=1,cols=4); set_widths(t,[3.45,0.75,1.05,1.2]); borders(t)
    heads=['Description','Qty','Rate','Amount']
    for c,h in zip(t.rows[0].cells,heads):
        shade(c,PALE); cell_margins(c); r=c.paragraphs[0].add_run(h); font(r,9.5,DARK,True)
    for _ in range(4):
        cells=t.add_row().cells
        for c in cells: cell_margins(c,150,120,150,120); r=c.paragraphs[0].add_run('[ ]'); font(r,10,NAVY)
    p(d,'Payment summary',style='Heading 1')
    summary=meta(d,[('Subtotal:','[0.00]'),('Discount:','[0.00]'),('Amount paid:','[0.00]'),('Total due:','[0.00]')])
    shade(summary.rows[-1].cells[0],PALE); shade(summary.rows[-1].cells[1],PALE)
    p(d,'Payment instructions',style='Heading 1')
    note(d,'Choose the applicable method:','CCP transfer: Account holder [name] | Account number [number] | RIP [number]. PayPal: [PayPal email or payment link]. Do not include payment details that you do not want shared with the recipient.')
    p(d,'Notes',style='Heading 1')
    p(d,'This invoice is governed by the signed project agreement and proposal. Please include the invoice number with your payment. The balance is due before final handover or production launch unless the proposal states otherwise.')
    p(d,CONTACT,color=MUTED,size=9.4)
    d.save(OUT/'Invoice_Template.docx')

if __name__=='__main__':
    onboarding(); agreement(); invoice()
